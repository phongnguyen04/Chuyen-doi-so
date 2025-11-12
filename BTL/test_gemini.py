import customtkinter as ctk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import google.generativeai as genai
import os
from dotenv import load_dotenv
import json
from docx import Document
from tkinter import filedialog
from PIL import Image



load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    messagebox.showerror("Lỗi", "Không tìm thấy GOOGLE_API_KEY trong file .env")
    exit()

genai.configure(api_key=api_key)
MODEL_NAME = "gemini-2.5-flash" 

def save_as_docx():
    text = result_box.get("1.0", "end").strip()
    if not text:
        messagebox.showwarning("Trống", "Không có nội dung để lưu!")
        return

    file_path = filedialog.asksaveasfilename(
        title="Lưu kết quả dưới dạng DOCX",
        defaultextension=".docx",
        filetypes=[("Microsoft Word", "*.docx")]
    )

    if file_path:
        try:
            doc = Document()
            doc.add_heading("Kết quả chỉnh sửa AI", level=1)
            doc.add_paragraph(text)
            doc.save(file_path)
            messagebox.showinfo("Thành công", "✅ Đã lưu file DOCX thành công!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu file DOCX:\n{e}")

# --- File lưu lịch sử ---
HISTORY_FILE = "history.json"

# --- Lưu và tải lịch sử ---
def save_history():
    """Lưu lịch sử ra file JSON"""
    try:
        with open(HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(history_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print("Lỗi khi lưu lịch sử:", e)

def load_history():
    """Tải lịch sử từ file JSON nếu có"""
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

# --- Dữ liệu lịch sử ---
history_data = load_history()

# --- Hàm hiển thị cửa sổ lịch sử ---
def open_history():
    if not history_data:
        messagebox.showinfo("Lịch sử trống", "Không có mục nào trong lịch sử.")
        return

    hist_win = ctk.CTkToplevel(app)
    hist_win.title("📜 Lịch sử chỉnh sửa")
    hist_win.geometry("720x520")
    hist_win.lift()
    hist_win.focus_force()
    hist_win.grab_set()




    # --- Khung chia 2 phần ---
    main_frame = ctk.CTkFrame(hist_win)
    main_frame.pack(fill="both", expand=True, padx=10, pady=10)

    # --- Bên trái: danh sách lịch sử ---
    listbox = ctk.CTkScrollableFrame(main_frame, width=230)
    listbox.pack(side="left", fill="y", padx=(0, 10))

    # --- Bên phải: khung chi tiết ---
    right_frame = ctk.CTkFrame(main_frame)
    right_frame.pack(side="right", fill="both", expand=True)

    detail_box = ctk.CTkTextbox(right_frame, wrap="word", font=("Consolas", 13))
    detail_box.pack(fill="both", expand=True, padx=5, pady=(5, 0))

    # --- Nút điều khiển ---
    btn_frame = ctk.CTkFrame(right_frame, fg_color="transparent")
    btn_frame.pack(pady=6)

    btn_delete = ctk.CTkButton(btn_frame, text="🗑️ Xóa mục này", width=140, fg_color="#E67E22")
    btn_delete.grid(row=0, column=0, padx=5)

    btn_clear_all = ctk.CTkButton(btn_frame, text="❌ Xóa toàn bộ", width=140, fg_color="#C21807")
    btn_clear_all.grid(row=0, column=1, padx=5)

    # --- Hàm hiển thị chi tiết ---
    def show_detail(index):
        item = history_data[index]
        detail_box.delete("1.0", "end")
        detail_box.insert("end", "📥 Gốc:\n" + item["input"] + "\n\n")
        detail_box.insert("end", "✅ Đã sửa:\n" + item["output"])
        # Gán chức năng xóa vào nút
        btn_delete.configure(command=lambda: delete_entry(index))

    # --- Hàm xóa 1 mục ---
    def delete_entry(index):
        if messagebox.askyesno("Xác nhận", "Bạn có chắc muốn xóa mục này?"):
            del history_data[index]
            save_history()
            hist_win.destroy()
            open_history()  # Mở lại giao diện mới

    # --- Hàm xóa toàn bộ ---
    def clear_all_history():
        if messagebox.askyesno("Xác nhận", "Bạn có chắc muốn xóa toàn bộ lịch sử?"):
            history_data.clear()
            save_history()
            hist_win.destroy()
            messagebox.showinfo("Đã xóa", "Toàn bộ lịch sử đã được xóa.")

    btn_clear_all.configure(command=clear_all_history)

    # --- Tạo nút cho từng mục trong lịch sử ---
    for i, item in enumerate(history_data):
        short_text = item['input'][:25].replace("\n", " ")
        btn = ctk.CTkButton(
            listbox,
            text=f"{i+1}. {short_text}...",
            width=200,
            anchor="w",
            command=lambda i=i: show_detail(i)
        )
        btn.pack(pady=2, fill="x")

    # --- Tự động chọn mục mới nhất ---
    show_detail(len(history_data)-1)


def upload_docx():
    file_path = filedialog.askopenfilename(
        title="Chọn file DOCX làm đầu vào",
        filetypes=[("Microsoft Word", "*.docx")]
    )

    if file_path:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            # Đưa nội dung vào ô nhập
            input_box.delete("1.0", "end")
            input_box.insert("end", text)

            # Nếu đang có ảnh → reset preview
            global image_path
            image_path = None
            preview_label.configure(image="", text="(Chưa có ảnh)")

            messagebox.showinfo("✅ Thành công", "Đã tải nội dung từ file DOCX!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file DOCX:\n{e}")
            
# --- Hàm xử lý văn bản hoặc ảnh ---
def check_spelling(input_text="", image_path=None):
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        prompt = (
            "Sửa lỗi chính tả, ngữ pháp và viết lại đoạn văn này bằng tiếng Việt tự nhiên chỉ viết lại câu sai sau đó chỉ ra lỗi sai và viết lại câu đúng.\n\n"
            "Nếu đây là ảnh, hãy đọc nội dung chữ viết trong ảnh, rồi sửa lỗi chính tả và viết lại cho đúng."
        )

        if image_path:
            image = Image.open(image_path)
            response = model.generate_content([prompt, image])
        else:
            response = model.generate_content(prompt + "\n\n" + input_text)

        return response.text
    except Exception as e:
        return f"❌ Lỗi khi xử lý: {e}"


def upload_image():
    global image_path, preview_img
    file_path = filedialog.askopenfilename(
        title="Chọn ảnh tay viết hoặc văn bản",
        filetypes=[("Ảnh", "*.png;*.jpg;*.jpeg;*.webp")]
    )
    if file_path:
        image_path = file_path
        input_box.delete("1.0", "end")

        img = Image.open(file_path)
        img.thumbnail((320, 180))
        preview_img = ImageTk.PhotoImage(img)
        preview_label.configure(image=preview_img, text="")
        messagebox.showinfo("Ảnh đã tải", "✅ Ảnh đã được chọn, ô nhập đã làm trống sẵn!")

def run_check():
    text_input = input_box.get("1.0", "end").strip()
    if not text_input and not image_path:
        messagebox.showwarning("Thiếu dữ liệu", "Hãy nhập đoạn văn hoặc tải ảnh lên!")
        return

    result_box.delete("1.0", "end")
    result_box.insert("end", "🧠 Đang xử lý... vui lòng chờ...\n")
    app.update_idletasks()

    result = check_spelling(text_input, image_path)

    result_box.delete("1.0", "end")
    result_box.insert("end", result)

    # ✅ Lưu vào lịch sử và file
    history_data.append({
        "input": text_input if text_input else "(Ảnh)",
        "output": result
    })
    save_history()

def clear_all():
    input_box.delete("1.0", "end")
    result_box.delete("1.0", "end")
    preview_label.configure(image="", text="(Chưa có ảnh)")
    global image_path
    image_path = None

def copy_result():
    text = result_box.get("1.0", "end").strip()
    app.clipboard_clear()
    app.clipboard_append(text)
    messagebox.showinfo("Đã copy", "📋 Kết quả đã được sao chép!")

# --- GUI chính ---
# 1. Cài đặt giao diện
ctk.set_appearance_mode("light")  # Chuyển sang giao diện Sáng
ctk.set_default_color_theme("blue")   # Dùng màu "blue" (hoặc "green") làm chủ đạo

# 2. Định nghĩa font chữ
# Font chữ chính cho các nhãn, nút bấm (sạch sẽ, hiện đại)
main_font = ("Segoe UI", 14) 
# Font chữ mono cho các ô nhập liệu (dễ đọc code/văn bản)
mono_font = ("Consolas", 14) 
 
app = ctk.CTk()
app.configure(fg_color="#EAF2F8") 
app.title("✨ Trình sửa lỗi chính tả (Chuyển đổi số)")
app.geometry("950x720")

# --- NẠP ICON (THÊM PHẦN NÀY) ---
# Tải ảnh bằng PIL và tạo đối tượng CTkImage
# Điều chỉnh size=(20, 20) cho phù hợp với bạn
try:
    icon_upload = ctk.CTkImage(Image.open("icons/upload_image.png"), size=(20, 20))
    icon_docx = ctk.CTkImage(Image.open("icons/upload_doc.png"), size=(20, 20))
    icon_check = ctk.CTkImage(Image.open("icons/check.png"), size=(20, 20))
    icon_clear = ctk.CTkImage(Image.open("icons/delete.png"), size=(20, 20))
    icon_copy = ctk.CTkImage(Image.open("icons/copy.png"), size=(20, 20))
    icon_export = ctk.CTkImage(Image.open("icons/save.png"), size=(20, 20))
    icon_history = ctk.CTkImage(Image.open("icons/history.png"), size=(20, 20))
    icon_exit = ctk.CTkImage(Image.open("icons/exit.png"), size=(20, 20))

except Exception as e:
    print(f"Lỗi khi nạp icon: {e}")
    # Đặt tất cả về None nếu có lỗi để app không bị crash
    icon_upload = icon_docx = icon_check = icon_clear = icon_copy = icon_export = icon_history = icon_exit = None

# --- Khung chính ---
frame = ctk.CTkFrame(app, corner_radius=15)
frame.pack(padx=20, pady=20, fill="both", expand=True)

# --- Tiêu đề & nhập văn bản ---
label_input = ctk.CTkLabel(
    frame, text="🖋️ Nhập đoạn văn cần sửa hoặc chọn ảnh bên dưới:",
    font=(main_font[0], 16, "bold") # Thay đổi ở đây
)
label_input.pack(pady=(10, 5))

input_box = ctk.CTkTextbox(frame, height=100, font=mono_font) # Thay đổi ở đây
input_box.pack(padx=10, pady=(0, 10), fill="x")

# --- Hiển thị ảnh đã chọn ---
# Dùng màu xám nhạt cho nền ảnh xem trước
preview_label = ctk.CTkLabel(frame, text="(Chưa có ảnh)", width=300, height=180, fg_color="#E0E0E0", corner_radius=10, text_color="#555555") # Thay đổi ở đây
preview_label.pack(pady=(0, 8))

image_path = None

# --- Kết quả ---
label_output = ctk.CTkLabel(frame, text="💎 Kết quả chỉnh sửa:", font=(main_font[0], 16, "bold")) # Thay đổi ở đây
label_output.pack(pady=(5, 5))

result_box = ctk.CTkTextbox(frame, height=220, font=mono_font) # Thay đổi ở đây
result_box.pack(padx=10, pady=(0, 5), fill="both", expand=True)
# --- Nút điều khiển ---
btn_frame = ctk.CTkFrame(frame, fg_color="transparent")
btn_frame.pack(pady=10, anchor="center")

# Định nghĩa màu sắc
COLOR_PRIMARY = "#3498DB" # Màu xanh dương chủ đạo (nếu theme là "blue")
# Nếu bạn dùng theme "green", hãy dùng màu này:
# COLOR_PRIMARY = "#2E8B57" 

COLOR_IO = "#27AE60"      # Xanh lá cho Tải lên/Xuất file
COLOR_NEUTRAL = "#7F8C8D"  # Xám cho các nút phụ
COLOR_WARNING = "#F39C12"  # Cam cho Cảnh báo (Xóa)
COLOR_DANGER = "#E74C3C"   # Đỏ cho Nguy hiểm (Thoát)

# Định nghĩa màu khi hover
COLOR_PRIMARY_HOVER = "#2980B9"
COLOR_IO_HOVER = "#229954"
COLOR_NEUTRAL_HOVER = "#707B7C"
COLOR_WARNING_HOVER = "#D68910"
COLOR_DANGER_HOVER = "#C0392B"


btn_upload = ctk.CTkButton(
    btn_frame, text="📷 Tải Ảnh", width=130, 
    fg_color=COLOR_IO, hover_color=COLOR_IO_HOVER, command=upload_image, font=main_font
)
btn_upload.grid(row=0, column=0, padx=6)

btn_import_docx = ctk.CTkButton(
    btn_frame, text="📄 Nhập DOCX", width=130, 
    fg_color=COLOR_IO, hover_color=COLOR_IO_HOVER, command=upload_docx, font=main_font
)
btn_import_docx.grid(row=0, column=1, padx=6)


btn_check = ctk.CTkButton(
    btn_frame, text="✨ Kiểm tra & Sửa lỗi", width=160, 
    fg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER, command=run_check, font=main_font
)
btn_check.grid(row=0, column=2, padx=6)

btn_clear = ctk.CTkButton(
    btn_frame, text="🧹 Xóa", width=100, 
    fg_color=COLOR_WARNING, hover_color=COLOR_WARNING_HOVER, command=clear_all, font=main_font
)
btn_clear.grid(row=0, column=3, padx=6)

btn_copy = ctk.CTkButton(
    btn_frame, text="📋 Copy", width=100, 
    fg_color=COLOR_NEUTRAL, hover_color=COLOR_NEUTRAL_HOVER, command=copy_result, font=main_font
)
btn_copy.grid(row=0, column=4, padx=6)

btn_export_docx = ctk.CTkButton(
    btn_frame, text="📄 Xuất DOCX", width=120, 
    fg_color=COLOR_IO, hover_color=COLOR_IO_HOVER, command=save_as_docx, font=main_font
)
btn_export_docx.grid(row=0, column=5, padx=6)

btn_history = ctk.CTkButton(
    btn_frame, text="📜 Lịch sử", width=100, 
    fg_color=COLOR_NEUTRAL, hover_color=COLOR_NEUTRAL_HOVER, command=open_history, font=main_font
)
btn_history.grid(row=0, column=6, padx=6)

btn_exit = ctk.CTkButton(
    btn_frame, text="❌ Thoát", width=100, 
    fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER, command=app.destroy, font=main_font
)
btn_exit.grid(row=0, column=7, padx=6)


status = ctk.CTkLabel(frame, text="⚡ Sẵn sàng – Vision Mode", font=("Segoe UI", 12))
status.pack(pady=(5, 8))
app.mainloop()
